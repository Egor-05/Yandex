from docx import Document


def markdown_to_docx(text):
    doc = Document()
    lns = text.split('\n')
    for line in lns:
        if line:
            # заголовки
            if line[:7].count('#') == 1:
                a = line.replace('#', '')
                doc.add_heading(a, level=1)
            elif line[:7].count('#') == 2:
                doc.add_heading(line.replace('#', ''), level=2)
            elif line[:7].count('#') == 3:
                doc.add_heading(line.replace('#', ''), level=3)
            elif line[:7].count('#') == 4:
                doc.add_heading(line.replace('#', ''), level=4)
            elif line[:7].count('#') == 5:
                doc.add_heading(line.replace('#', ''), level=5)
            elif line[:7].count('#') == 6:
                doc.add_heading(line.replace('#', ''), level=6)
            # списки
            elif str(line[:2]) == '- ':
                doc.add_paragraph(line[2:], style='List Bullet')
            elif str(line[:2]) == '* ':
                doc.add_paragraph(line[2:], style='List Bullet')
            elif str(line[:2]) == '+ ':
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line[0].isdigit() and line[1] == '.':
                doc.add_paragraph(line[2:], style='List Number')
            # различные выделения букв
            elif line[:3].count('*') == 1:
                doc.add_paragraph().add_run(line[1:-1]).italic = True
            elif line[:3].count('*') == 2:
                doc.add_paragraph().add_run(line[2:-2]).bold = True
            elif line[:3].count('*') == 3:
                doc.add_paragraph().add_run(line[3:-3]).italic = True
                doc.add_paragraph().add_run(line[3:-3]).bold = True
            # строки не содержащие ничеко из выше перечисленного
            else:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph()
    doc.save('res.docx')


markdown_to_docx('test01\nАбзацы создаются при помощи пустой строки. Если вокруг текста сверху и снизу есть пустые строки, то текст превращается в абзац.')
            #'\n\nЧтобы сделать перенос строки вместо абзаца,\nнужно поставить два пробела в конце предыдущей строки.\n\nЗаголовки отмечаются диезом `#` в начале строки, от одного до шести. Например:\n\n# Заголовок первого уровня\n## Заголовок h2\n### Заголовок h3\n#### Заголовок h4\n##### Заголовок h5\n###### Заголовок h6\n\nВ декоративных целях заголовки можно «закрывать» с обратной стороны.\n\n### Списки\n\nДля разметки неупорядоченных списков можно использовать или `*`, или `-')


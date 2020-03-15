from sys import stdin
from docx import Document


place = input().lower()
time = input().lower()
names = [''.join([j for j in i if j.isalpha() or j == ' ']).strip() for i in stdin]
document = Document()
for i in names:

    document.add_heading('Приглашение', 0)

    p = document.add_paragraph(f'Здравствуйте {i}, приглашаем Вас на мероприятие в'
                               f' честь 8 марта, которое состоится в')
    p.add_run(f'{place} {time}')
    if names[names.index(i)] != names[-1]:
        document.add_page_break()
document.save('test.docx')
